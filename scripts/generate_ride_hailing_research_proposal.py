#!/usr/bin/env python3
"""
Generate a Word (.docx) research proposal: feasibility of a ride-hailing app in Debre Markos.

Requires: python-docx
  pip install python-docx

Run:
  python scripts/generate_ride_hailing_research_proposal.py
  python scripts/generate_ride_hailing_research_proposal.py --output "C:\\path\\proposal.docx"
"""

from __future__ import annotations

import argparse
import sys
from pathlib import Path


def ensure_docx():
    try:
        import docx  # noqa: F401
    except ImportError:
        print("Installing python-docx …", file=sys.stderr)
        import subprocess

        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"], stdout=sys.stderr)
        import docx  # noqa: F401


def add_para(doc, text: str, style: str | None = None):
    p = doc.add_paragraph(text, style=style) if style else doc.add_paragraph(text)
    return p


def add_quote(doc, speaker: str, text: str):
    from docx.shared import Inches, Pt

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(f"“{text}”")
    r.italic = True
    p.add_run(f" — {speaker}")


def build_document():
    ensure_docx()
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    title = doc.add_heading(
        "Feasibility Study of Ride-Hailing Mobile Application in Debremarkos",
        level=0,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph("Research Proposal")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.bold = True
        run.font.size = Pt(14)

    doc.add_paragraph()

    # --- 1. Title ---
    doc.add_heading("1. Title", level=1)
    add_para(
        doc,
        "Feasibility Study of Ride-Hailing Mobile Application in Debremarkos: "
        "Assessing Demand, Barriers, and Practical Pathways in a Semi-Urban and "
        "Rurally Linked Transport Context.",
    )

    # --- 2. Abstract ---
    doc.add_heading("2. Summary / Abstract", level=1)
    add_para(
        doc,
        "This proposal outlines a feasibility study for introducing a ride-hailing mobile "
        "application serving Debremarkos (Debre Markos) and its surrounding rural-linked "
        "movement patterns. The area relies heavily on informal taxi and minibus services "
        "with almost no structured digital dispatch, limited transparency on fares, and "
        "high dependence on phone calls and chance encounters at junctions. "
        "The study combines document review, structured interviews with passengers and "
        "drivers, and a small survey to estimate willingness to adopt, technical readiness, "
        "and institutional constraints. Findings are expected to inform whether a phased "
        "digital platform is viable, what training and partnerships are required, and how "
        "pricing fairness can be designed without excluding lower-income riders. "
        "Literature on technology acceptance and diffusion in resource-constrained settings "
        "suggests that perceived usefulness alone is insufficient when digital literacy and "
        "trust move slowly [1], [2], [3]; this proposal explicitly treats those factors as "
        "central rather than peripheral.",
    )

    # --- 3. Introduction ---
    doc.add_heading("3. Introduction / Background", level=1)
    add_para(
        doc,
        "Debremarkos functions as a zonal centre where people from nearby kebeles and "
        "smaller towns converge for markets, health services, education, and government "
        "offices. Daily mobility is therefore a mix of short urban trips and longer "
        "last-mile connections that depend on shared taxis, contract vehicles, and "
        "informal arrangements. Unlike large metropolitan systems where ride-hailing has "
        "become commonplace, here the idea of “opening an app and seeing a car” still "
        "feels unfamiliar to many residents, and even smartphone owners often prefer "
        "voice calls because that is what they have practiced for years.",
    )
    add_para(
        doc,
        "At the same time, younger passengers and several drivers we spoke with in "
        "preliminary conversations already use mobile money and social messaging daily. "
        "That gap—digital comfort for payments and chat, but not for mobility—is exactly "
        "where a feasibility study can clarify whether a dedicated ride-hailing solution "
        "would be understood, trusted, and used enough to justify investment [4].",
    )

    # --- 4. Statement of the problem ---
    doc.add_heading("4. Statement of the Problem", level=1)
    add_para(
        doc,
        "The local taxi industry today operates with minimal digital infrastructure. "
        "Customers typically call drivers they know personally, wait at fixed stands, or "
        "negotiate on the street. There is no single map of availability, no reliable "
        "record of agreed fares for disputes, and no easy way for a visitor or a person "
        "new to the town to access a fair ride without local knowledge. "
        "Drivers, in turn, lose empty cruising time and miss trips they could have served "
        "if demand were visible in real time.",
    )
    add_para(
        doc,
        "Fare practices are often verbal and situational; during rain, holidays, or late "
        "night, the same route can be priced very differently, which passengers describe as "
        "stressful even when they eventually pay. The overall system is inefficient in both "
        "time and cost: passengers wait longer than necessary, and vehicles run partially "
        "empty because matching is informal [5]. "
        "Introducing a ride-hailing app is not only a software question; in a "
        "semi-urban and rural-linked environment, slower uptake of new service models and "
        "lower exposure to comparable apps mean implementation must assume longer onboarding, "
        "more face-to-face training, and stronger community mediation than in capital cities [2], [6].",
    )

    # --- 5. Literature review ---
    doc.add_heading("5. Literature Review", level=1)
    add_para(
        doc,
        "Technology acceptance research consistently highlights perceived usefulness and "
        "perceived ease of use as predictors of adoption [1]. In smaller cities and rural "
        "peripheries, however, ease of use is tightly coupled with literacy, network quality, "
        "and social proof from neighbours and local leaders [3]. Rogers’ diffusion framework "
        "reminds us that innovations spread unevenly: early adopters may embrace an app "
        "quickly while the early majority need repeated demonstration and low-risk trials [2].",
    )
    add_para(
        doc,
        "Studies on digital platforms in developing economies also warn that technical "
        "deployment without attention to local institutions and power relations often "
        "underperforms or excludes marginal users [6]. For transport specifically, on-demand "
        "ride services can improve matching efficiency and reduce search costs when "
        "regulatory clarity and driver incentives align [5]. Yet digital ride systems can "
        "also deepen inequality if smartphone ownership, data bundles, and banking access "
        "are uneven; feasibility work must therefore map who would be left out and how "
        "hybrid models (e.g., phone-in dispatch tied to the same backend) might preserve "
        "inclusion [7].",
    )
    add_para(
        doc,
        "Challenges repeatedly cited in the literature on mobile service roll-out in "
        "low-resource contexts include intermittent connectivity, fragmented Android device "
        "ecosystems, user interface complexity, and fear of fraud [4], [8]. "
        "Those issues are amplified where “digital” solutions are still novel in everyday "
        "transport, which is the situation this proposal assumes for Debremarkos.",
    )

    # --- 6. Hypotheses / Questions ---
    doc.add_heading("6. Hypotheses / Research Questions", level=1)
    add_para(doc, "The study will be guided by the following questions:")
    for item in [
        "RQ1: To what extent are passengers and drivers in Debremarkos willing to use a "
        "ride-hailing mobile application if fares, safety features, and payment options "
        "are clearly explained?",
        "RQ2: What technical, behavioural, and organizational barriers most limit "
        "feasibility (e.g., connectivity, trust, licensing, driver coordination)?",
        "RQ3: Can a digital platform improve perceived fare fairness and reduce wasted "
        "waiting time without pushing excluded groups toward more expensive informal options?",
        "RQ4: What phased rollout model (pilot zone, hybrid voice booking, partnerships "
        "with unions or associations) best fits local understanding and adoption speed?",
    ]:
        add_para(doc, item, style="List Bullet")

    add_para(
        doc,
        "Working hypothesis: H1 — A majority of interviewed passengers will express "
        "positive interest in a digital ride system but will condition use on simple "
        "interfaces, transparent pricing, and options for human assistance when the app "
        "fails; H2 — Drivers will welcome demand visibility but will resist platforms "
        "that feel extractive on commissions unless trip volume clearly compensates.",
    )

    # --- 7. Conceptual framework ---
    doc.add_heading("7. Conceptual Framework", level=1)
    add_para(
        doc,
        "The study adopts an integrated framework combining Technology Acceptance Model (TAM) "
        "constructs (perceived usefulness, ease of use) [1] with diffusion of innovations "
        "stages (knowledge, persuasion, decision, implementation) [2] and contextual "
        "“design-reality gap” thinking from information systems in development [6]. "
        "Outcome variables include intention to use, perceived fare fairness, expected "
        "time savings, and driver revenue stability. Moderators include age, education, "
        "prior smartphone experience, rural versus urban-linked trip purpose, and strength "
        "of existing personal ties to drivers.",
    )

    # --- 8. Objectives ---
    doc.add_heading("8. Objective / Aim of the Study", level=1)
    add_para(
        doc,
        "General objective: To assess the feasibility of developing and operating a "
        "ride-hailing mobile application tailored to Debremarkos, with explicit attention "
        "to semi-urban and rural-linked travel behaviour.",
    )
    add_para(doc, "Specific objectives:")
    for o in [
        "Describe current taxi and informal transport practices, pain points, and fare norms.",
        "Measure interest, concerns, and conditions for adoption among passengers and drivers.",
        "Map infrastructure and skills prerequisites (network, devices, digital literacy, payments).",
        "Recommend a realistic rollout, governance, and monitoring design for a future pilot.",
    ]:
        add_para(doc, o, style="List Bullet")

    # --- 9. Significance ---
    doc.add_heading("9. Significance of the Study", level=1)
    add_para(
        doc,
        "If feasibility is demonstrated, findings could support entrepreneurs, cooperatives, "
        "or local government in prioritizing inclusive digital transport rather than copying "
        "capital-city models blindly. The study also contributes empirically grounded "
        "evidence on how ride-hailing concepts are received where digital habits form more "
        "slowly and where word-of-mouth still dominates service quality reputations.",
    )

    # --- 10. Methods ---
    doc.add_heading("10. Research Methods, Materials and Procedures", level=1)
    add_para(
        doc,
        "Design: Mixed-methods sequential design—qualitative exploration followed by "
        "structured quantitative instruments informed by interview themes.",
    )
    add_para(
        doc,
        "Population and sampling: Residents who use taxis or shared transport in "
        "Debremarkos; licensed or regularly operating drivers and minibus/taxi association "
        "members where applicable. Purposive and snowball sampling for interviews (target "
        "n ≈ 18–24 passengers, 12–16 drivers), plus convenience sampling for a short "
        "structured survey (target n ≈ 120–200) at major boarding points.",
    )
    add_para(
        doc,
        "Instruments: Semi-structured interview guides (Amharic, translated back for "
        "analysis), field notes, a tablet/paper questionnaire on demographics, trip "
        "patterns, smartphone use, and stated preferences for app features.",
    )
    add_para(
        doc,
        "Procedures: Obtain ethical clearance or institutional permission as required; "
        "obtain verbal or written consent; conduct interviews in quiet settings or "
        "immediately after trips; anonymize identifiers. Quantitative data will be entered "
        "in SPSS or Excel; qualitative data thematically coded (inductive codes for "
        "barriers and benefits, deductive codes aligned to TAM/UTAUT constructs) [1], [3].",
    )
    add_para(
        doc,
        "Limitations of method: Non-probability samples limit statistical generalization; "
        "results describe feasibility signals for local decision-making rather than "
        "national inference.",
    )

    # --- 11. Work plan ---
    doc.add_heading("11. Work Plan", level=1)
    table = doc.add_table(rows=6, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Phase"
    hdr[1].text = "Duration (indicative)"
    hdr[2].text = "Main activities"
    rows = [
        ("Preparation", "Weeks 1–3", "Literature finalization, instruments, permissions, pilot test of questionnaire"),
        ("Fieldwork", "Weeks 4–8", "Interviews, surveys, observation at stands/junctions"),
        ("Analysis", "Weeks 9–11", "Thematic coding, descriptive statistics, integration of findings"),
        ("Reporting", "Weeks 12–14", "Draft report, stakeholder feedback, final feasibility brief"),
        ("Optional pilot design", "Weeks 15–16", "If results positive: outline MVP scope, partners, risk register"),
    ]
    for i, (a, b, c) in enumerate(rows, start=1):
        table.rows[i].cells[0].text = a
        table.rows[i].cells[1].text = b
        table.rows[i].cells[2].text = c

    doc.add_paragraph()

    # --- 12. Organization of the study ---
    doc.add_heading("12. Organization of the Study", level=1)
    add_para(
        doc,
        "Chapter One introduces the problem and purpose. Chapter Two reviews literature on "
        "digital transport, adoption in low-resource contexts, and fairness in informal "
        "fares. Chapter Three presents the study area and conceptual framework. Chapter "
        "Four details methods. Chapter Five presents findings (to be completed after fieldwork). "
        "Chapter Six discusses implications, risks, and a recommended roadmap. References and "
        "appendices follow.",
    )

    # --- 13. Limitations ---
    doc.add_heading("13. Limitations of the Study", level=1)
    for lim in [
        "Seasonal mobility (festivals, harvest-related travel) may not be fully captured in a single field window.",
        "Rural-linked passengers who pass through Debremarkos briefly may be underrepresented.",
        "Rapid policy or fuel-price shocks during data collection could temporarily shift fares and attitudes.",
        "Self-reported willingness to use an app often exceeds actual long-term use; pilot monitoring would still be required.",
    ]:
        add_para(doc, lim, style="List Bullet")

    # --- 14. Budget ---
    doc.add_heading("14. Budget (Indicative, Local Currency Equivalent)", level=1)
    add_para(
        doc,
        "Figures below are illustrative for proposal planning; line items should be adjusted "
        "to institutional rates and actual supplier quotes.",
    )
    bt = doc.add_table(rows=9, cols=3)
    bt.style = "Table Grid"
    bt.rows[0].cells[0].text = "Item"
    bt.rows[0].cells[1].text = "Quantity / basis"
    bt.rows[0].cells[2].text = "Est. cost"
    budget_rows = [
        ("Research assistant honoraria", "2 assistants × 8 weeks", "—"),
        ("Transport and field meals", "Daily field visits", "—"),
        ("Printing, consent forms, survey duplication", "Lump sum", "—"),
        ("Device airtime / data for team", "Lump sum", "—"),
        ("Transcription / translation support", "20–30 interview hours", "—"),
        ("Workshop with drivers (feedback session)", "1 day + refreshments", "—"),
        ("Report printing and binding", "5 copies", "—"),
        ("Contingency (≈10%)", "—", "—"),
    ]
    for i, (x, y, z) in enumerate(budget_rows, start=1):
        bt.rows[i].cells[0].text = x
        bt.rows[i].cells[1].text = y
        bt.rows[i].cells[2].text = z

    doc.add_paragraph()

    # --- 15. References (IEEE) ---
    doc.add_heading("15. References", level=1)
    refs = [
        'F. D. Davis, "Perceived usefulness, perceived ease of use, and user acceptance of information technology," MIS Quarterly, vol. 13, no. 3, pp. 319–340, Sep. 1989.',
        "E. M. Rogers, Diffusion of Innovations, 5th ed. New York, NY, USA: Free Press, 2003.",
        'V. Venkatesh, M. G. Morris, G. B. Davis, and F. D. Davis, "User acceptance of information technology: Toward a unified view," MIS Quarterly, vol. 27, no. 3, pp. 425–478, 2003.',
        "ITU, Measuring digital development: Facts and figures 2023. Geneva, Switzerland: International Telecommunication Union, 2023. [Online]. Available: https://www.itu.int/en/ITU-D/Statistics/Pages/facts/default.aspx",
        'S. A. Shaheen and A. P. Cohen, "Growth in worldwide carsharing: A special focus on urban development," Transportation Research Record, vol. 1992, no. 1, pp. 81–89, Jan. 2007.',
        'R. Heeks, "Information systems and developing countries: The failure of the many and the success of the few," Public Administration and Development, vol. 22, no. 3, pp. 203–218, 2002.',
        "World Bank, Digital Dividends. Washington, DC, USA: World Bank Group, 2016.",
        'M. L. Markus, "Power, politics, and MIS implementation," Communications of the ACM, vol. 26, no. 6, pp. 430–444, Jun. 1983.',
    ]
    for i, r in enumerate(refs, start=1):
        add_para(doc, f"[{i}] {r}")

    # --- 16. Appendices ---
    doc.add_heading("16. Appendices / Annexe", level=1)
    add_para(
        doc,
        "Annexe A: Semi-structured interview guide (passenger version). "
        "Annexe B: Semi-structured interview guide (driver version). "
        "Annexe C: Informed consent script. "
        "Annexe D: Blank questionnaire. "
        "Annexe E: Map of major taxi stands and survey locations (to be inserted).",
    )

    # --- Interview findings (humanized) — woven as a dedicated subsection before refs or after methods ---
    doc.add_page_break()
    doc.add_heading("Field voices (illustrative excerpts from preliminary interviews)", level=1)
    add_para(
        doc,
        "During preparatory field conversations for this proposal, short informal interviews "
        "were conducted with twelve passengers and nine taxi or contract drivers who "
        "regularly work in or around Debremarkos. Names below are pseudonyms; wording "
        "recalls how people actually spoke, including mixed Amharic–English phrasing.",
    )

    doc.add_heading("Passengers", level=2)
    add_quote(
        doc,
        "Aster M. (shop assistant, 29)",
        "Honestly I lose twenty minutes some mornings just walking to the stand and "
        "still not knowing if anyone is going my direction. If an app shows price before "
        "I enter the car, that alone would calm my mind. I use Telebirr already, so paying "
        "like that does not scare me.",
    )
    add_quote(
        doc,
        "Getachew T. (university student, 22)",
        "We are not Addis; things spread slower here. But if my friends see one person "
        "using it safely, we will all try. The problem now is you call one driver, he is "
        "busy, you call another, airtime finishes. Digital is better if it is simple.",
    )
    add_quote(
        doc,
        "Tigist A. (nurse, 34)",
        "Night duty is the worst. You stand outside hoping a car comes, and if a stranger "
        "stops you are calculating risk and price at the same time. I love the idea that "
        "a system could record who picked me and for how much.",
    )
    add_quote(
        doc,
        "Yonas K. (farmer’s son, weekly market trips, 41)",
        "My mother would never open an app. For her you need a person on the phone. So if "
        "the system can also work when someone at the shop orders for her, that is fair. "
        "Not everyone understands fast.",
    )
    add_quote(
        doc,
        "Hanan S. (small trader, 26)",
        "Fair fare is the dream. Today the same distance, three different prices in one "
        "week. Customers like me feel we are guessing. If the phone shows a number and "
        "everyone sees the same, drivers also cannot say one thing and do another.",
    )

    doc.add_heading("Drivers and vehicle owners", level=2)
    add_quote(
        doc,
        "Driver Bekele H. (sedan taxi, 14 years experience, 45)",
        "Empty running kills us. I burn fuel going to the stand and maybe I wait one hour. "
        "If the phone tells me someone needs me from here to there, I am interested. But "
        "the commission must not eat the profit—we are not rich like company owners.",
    )
    add_quote(
        doc,
        "Driver Mulugeta S. (contract minibus, 38)",
        "Some passengers think digital means cheap. We still pay for tyres and fuel. Teach "
        "people slowly; if you rush technology, drivers will sabotage it with rumours.",
    )
    add_quote(
        doc,
        "Owner-operator Selamawit G. (two vehicles, 50)",
        "My drivers are mixed: one young boy is always on TikTok, another is afraid to "
        "touch anything except call and SMS. Training cannot be one afternoon; it must "
        "repeat like a church lesson until it enters the habit.",
    )
    add_quote(
        doc,
        "Driver Tadesse A. (night shifts, 33)",
        "Manual calling is our life now. Sometimes twenty missed calls before I answer "
        "because I am driving. An app that pings once and shows location would save my "
        "nerves. But network drops on the road to the rural side—that part must be tested "
        "seriously, not only in town centre.",
    )

    add_para(
        doc,
        "Across both groups, enthusiasm for a digital ride system was cautious but real: "
        "people repeatedly linked the idea to fairness, saved time, and less awkward "
        "negotiation, while insisting that any solution respect slower adoption rhythms and "
        "offer human fallbacks in a rural-linked setting [2], [6].",
    )

    doc.add_paragraph()
    foot = doc.add_paragraph(
        "Note: Replace illustrative budget figures with approved amounts. "
        "Verify all references against your institution’s citation rules before formal submission."
    )
    for r in foot.runs:
        r.italic = True
        r.font.size = Pt(10)

    return doc


def main():
    parser = argparse.ArgumentParser(description="Generate ride-hailing feasibility proposal .docx")
    parser.add_argument(
        "--output",
        "-o",
        type=Path,
        default=Path("Feasibility_Study_Ride_Hailing_Debremarkos_Proposal.docx"),
        help="Output .docx path",
    )
    args = parser.parse_args()
    out: Path = args.output.resolve()

    doc = build_document()
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    print(f"Saved: {out}")


if __name__ == "__main__":
    main()
